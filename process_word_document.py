import uuid
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from minio import Minio
from minio.error import S3Error
import docx.opc.constants
import docx.oxml
from docx.oxml.shared import qn
import os
from pathlib import Path
# MinIO 配置
MINIO_HOST = '172.16.0.133:9000'
MINIO_ACCESS_KEY = 'rag_flow'
MINIO_SECRET_KEY = 'infini_rag_flow'
BUCKET_NAME = 'ragflow-imges'
# 定义 pic 和 a 命名空间
PIC_NS = {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}
DRAWINGML_NS = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
def upload_image_to_minio(image_data, image_ext):
    """优化图片类型检测和上传"""
    try:
        minio_client = Minio(
            MINIO_HOST,
            access_key=MINIO_ACCESS_KEY,
            secret_key=MINIO_SECRET_KEY,
            secure=False
        )
        if not minio_client.bucket_exists(BUCKET_NAME):
            minio_client.make_bucket(BUCKET_NAME)
            print(f"Created bucket '{BUCKET_NAME}'")
        # 根据文件头确定扩展名和类型
        content_types = {
            b'\xFF\xD8\xFF': ('jpg', 'image/jpeg'),
            b'\x89PNG': ('png', 'image/png')
        }
        for sig, (ext, ctype) in content_types.items():
            if image_data.startswith(sig):
                image_ext = ext
                content_type = ctype
                break
        else:  # 默认jpg
            image_ext = 'jpg'
            content_type = 'image/jpeg'
        image_name = f"{uuid.uuid4()}.{image_ext}"
        image_url = f"![image](/{BUCKET_NAME}/{image_name})"
        minio_client.put_object(
            BUCKET_NAME,
            image_name,
            io.BytesIO(image_data),
            len(image_data),
            content_type=content_type
        )
        print(f"Image uploaded: {image_url}")
        return image_url
    except (S3Error, Exception) as e:
        print(f"Image upload failed: {e}")
        return None
def create_hyperlink(paragraph, url):
    """创建文本链接并居中"""
    run = paragraph.add_run(f"{url}") # 输出图片链接
    #paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  #居中
    return run
def process_word_document(file_path, output_dir):
    """优化路径处理和图片提取逻辑"""
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        return None
    try:
        doc = Document(file_path)
        img_count = 0
        # 处理段落内图片
        for para in doc.paragraphs:
            for inline in para.runs:
                # 新版docx图片提取方式
                img = inline._element.find('.//pic:pic', namespaces=PIC_NS)  # 使用定义的命名空间
                if img is None:
                    continue
                blip = img.find('.//a:blip', namespaces=DRAWINGML_NS) # 使用DRAWINGML_NS
                if blip is None:
                    continue
                r_embed = blip.attrib.get(qn('r:embed'))
                part = doc.part.related_parts[r_embed]
                img_data = part.blob
                if not img_data:
                    continue
                img_url = upload_image_to_minio(img_data, "")
                if img_url:
                    # 清除原图片元素
                    for c in list(inline._element):
                        if c.tag.endswith('drawing'):
                            inline._element.remove(c)
                    create_hyperlink(para, img_url)
                    img_count += 1
        # 处理表格图片
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # 复用段落内图片处理逻辑
                        for inline in para.runs:
                            img = inline._element.find('.//pic:pic', namespaces=PIC_NS)  # 使用定义的命名空间
                            if img is None:
                                continue
                            blip = img.find('.//a:blip', namespaces=DRAWINGML_NS) # 使用DRAWINGML_NS
                            if blip is None:
                                continue
                            r_embed = blip.attrib.get(qn('r:embed'))
                            part = doc.part.related_parts[r_embed]
                            img_data = part.blob
                            if not img_data:
                                continue
                            img_url = upload_image_to_minio(img_data, "")
                            if img_url:
                                # 清除原图片元素
                                for c in list(inline._element):
                                    if c.tag.endswith('drawing'):
                                        inline._element.remove(c)
                                create_hyperlink(para, img_url)
                                img_count += 1
        # 输出修改后的文件
        os.makedirs(output_dir, exist_ok=True)  # 确保输出目录存在
        filename = os.path.basename(file_path)
        new_path = os.path.join(output_dir, filename)  # 使用原始文件名
        doc.save(new_path)
        print(f"文件处理完成: {new_path}")
        return new_path
    except Exception as e:
        print(f"处理文件 {file_path} 时发生错误: {e}")
        return None
def process_directory(dir_path, output_dir):
    """处理指定目录下的所有 Word 文档"""
    if not os.path.exists(dir_path):
        print(f"目录不存在: {dir_path}")
        return
    for filename in os.listdir(dir_path):
        if filename.endswith(".docx") or filename.endswith(".doc"):
            file_path = os.path.join(dir_path, filename)
            process_word_document(file_path, output_dir)
if __name__ == '__main__':
    input_directory = "input_documents"  #  输入文档的目录
    output_directory = "output_documents"  #  输出目录
    # 确保输入目录存在
    if not os.path.exists(input_directory):
        print(f"创建输入目录: {input_directory}")
        os.makedirs(input_directory)
    #  处理目录中的所有 Word 文档
    process_directory(input_directory, output_directory)
